# compatibility/word_compatibility.py
"""Microsoft Word compatibility module for The Edison."""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
import logging

logger = logging.getLogger(__name__)

class WordCompatibility:
    """
    Ensures compatibility with Microsoft Word.
    Handles Word-specific formatting requirements and quirks.
    """
    
    def __init__(self):
        """Initialize Word compatibility handler."""
        self.document = None
    
    def create_document(self):
        """
        Create a new Word document.
        
        Returns:
            Document: A python-docx Document object
        """
        self.document = Document()
        return self.document
    
    def apply_mla_format(self, document, metadata):
        """
        Apply MLA formatting to a Word document.
        
        Args:
            document: python-docx Document object
            metadata (dict): Document metadata
            
        Returns:
            Document: Formatted document
        """
        # Set 1-inch margins
        for section in document.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Set font to Times New Roman, 12pt
        style = document.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        
        # Set double spacing
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        
        # Create header with student info
        if 'author' in metadata:
            p = document.add_paragraph(metadata.get('author', ''))
            document.add_paragraph(metadata.get('instructor', ''))
            document.add_paragraph(metadata.get('course', ''))
            document.add_paragraph(metadata.get('date', ''))
            
            # Add title (centered)
            if 'title' in metadata:
                title = document.add_paragraph(metadata['title'])
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add a blank paragraph after title
            document.add_paragraph()
        
        # Add header with page numbers
        header = document.sections[0].header
        header_para = header.paragraphs[0]
        header_para.text = f"{metadata.get('author', '')} "
        header_para.add_run().add_field('PAGE')
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        return document
    
    def apply_apa_format(self, document, metadata):
        """
        Apply APA formatting to a Word document.
        
        Args:
            document: python-docx Document object
            metadata (dict): Document metadata
            
        Returns:
            Document: Formatted document
        """
        # Set 1-inch margins
        for section in document.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Set font to Times New Roman or Arial, 12pt
        style = document.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        
        # Set double spacing
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        
        # Create title page
        if 'title' in metadata:
            # Title centered, bold, upper half of page
            title = document.add_paragraph()
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title.add_run(metadata['title'])
            title_run.bold = True
            
            # Author
            if 'author' in metadata:
                author = document.add_paragraph(metadata['author'])
                author.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Institution
            if 'institution' in metadata:
                institution = document.add_paragraph(metadata['institution'])
                institution.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
            # Course, instructor, date
            course_info = document.add_paragraph()
            course_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if 'course' in metadata:
                course_info.add_run(metadata['course'] + '\n')
            if 'instructor' in metadata:
                course_info.add_run(metadata['instructor'] + '\n')
            if 'date' in metadata:
                course_info.add_run(metadata['date'])
            
            # Add section break
            document.add_section(WD_SECTION.NEW_PAGE)
        
        # Add running head
        header = document.sections[0].header
        header_para = header.paragraphs[0]
        header_para.text = f"Running head: {metadata.get('title', '').upper()[:50]}"
        header_para.add_run("\t").add_field('PAGE')
        
        return document
    
    def apply_chicago_format(self, document, metadata):
        """
        Apply Chicago formatting to a Word document.
        
        Args:
            document: python-docx Document object
            metadata (dict): Document metadata
            
        Returns:
            Document: Formatted document
        """
        # Set 1-inch margins
        for section in document.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Set font to Times New Roman, 12pt
        style = document.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        
        # Set double spacing
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        
        # Create title page
        if 'title' in metadata:
            # Title centered, about 1/3 down the page
            for _ in range(10):  # Add space to move title down
                document.add_paragraph()
                
            title = document.add_paragraph()
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title.add_run(metadata['title'])
            
            # Add some space
            for _ in range(2):
                document.add_paragraph()
            
            # Author
            if 'author' in metadata:
                author = document.add_paragraph(metadata['author'])
                author.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add more space
            for _ in range(10):
                document.add_paragraph()
            
            # Course, instructor, date at bottom
            bottom_info = document.add_paragraph()
            bottom_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if 'course' in metadata:
                bottom_info.add_run(metadata['course'] + '\n')
            if 'instructor' in metadata:
                bottom_info.add_run(metadata['instructor'] + '\n')
            if 'date' in metadata:
                bottom_info.add_run(metadata['date'])
            
            # Add section break
            document.add_section(WD_SECTION.NEW_PAGE)
        
        # Add page numbers
        # Chicago style has page number at bottom center for first pages of chapters
        # and top right for subsequent pages
        first_page_section = document.sections[0]
        first_page_footer = first_page_section.footer
        footer_para = first_page_footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_para.add_run().add_field('PAGE')
        
        # For other sections/pages
        if len(document.sections) > 1:
            for section in document.sections[1:]:
                header = section.header
                header_para = header.paragraphs[0]
                header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                header_para.add_run().add_field('PAGE')
        
        return document
    
    def format_citations_mla(self, document, citations):
        """
        Format citations in MLA style for Word.
        
        Args:
            document: python-docx Document object
            citations (list): Citations to format
            
        Returns:
            Document: Document with formatted citations
        """
        # Add a works cited page
        document.add_section(WD_SECTION.NEW_PAGE)
        
        # Add "Works Cited" header, centered
        works_cited = document.add_paragraph("Works Cited")
        works_cited.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add each citation as a paragraph with hanging indent
        for citation in citations:
            p = document.add_paragraph(citation)
            p.paragraph_format.first_line_indent = Pt(-36)  # Negative for hanging indent
            p.paragraph_format.left_indent = Pt(36)
        
        return document
    
    def format_citations_apa(self, document, citations):
        """
        Format citations in APA style for Word.
        
        Args:
            document: python-docx Document object
            citations (list): Citations to format
            
        Returns:
            Document: Document with formatted citations
        """
        # Add a references page
        document.add_section(WD_SECTION.NEW_PAGE)
        
        # Add "References" header, centered
        references = document.add_paragraph("References")
        references.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add each citation as a paragraph with hanging indent
        for citation in citations:
            p = document.add_paragraph(citation)
            p.paragraph_format.first_line_indent = Pt(-36)  # Negative for hanging indent
            p.paragraph_format.left_indent = Pt(36)
        
        return document
    
    def format_citations_chicago(self, document, citations, notes=None):
        """
        Format citations in Chicago style for Word.
        
        Args:
            document: python-docx Document object
            citations (list): Citations to format
            notes (list): Footnotes to format
            
        Returns:
            Document: Document with formatted citations
        """
        # Add footnotes if provided
        if notes:
            for i, note in enumerate(notes, 1):
                document.add_footnote(note)
        
        # Add a bibliography page
        document.add_section(WD_SECTION.NEW_PAGE)
        
        # Add "Bibliography" header, centered
        bibliography = document.add_paragraph("Bibliography")
        bibliography.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add each citation as a paragraph with hanging indent
        for citation in citations:
            p = document.add_paragraph(citation)
            p.paragraph_format.first_line_indent = Pt(-36)  # Negative for hanging indent
            p.paragraph_format.left_indent = Pt(36)
        
        return document


# compatibility/google_docs_compatibility.py
"""Google Docs compatibility module for The Edison."""

import logging
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
from google.oauth2.credentials import Credentials

logger = logging.getLogger(__name__)

class GoogleDocsCompatibility:
    """
    Ensures compatibility with Google Docs.
    Handles Google Docs-specific formatting requirements and API integration.
    """
    
    def __init__(self, credentials=None):
        """
        Initialize Google Docs compatibility handler.
        
        Args:
            credentials: Google API credentials
        """
        self.credentials = credentials
        self.docs_service = None
        self.drive_service = None
    
    def authenticate(self, credentials):
        """
        Authenticate with Google API.
        
        Args:
            credentials: Google API credentials
            
        Returns:
            bool: True if authentication successful, False otherwise
        """
        try:
            self.credentials = credentials
            self.docs_service = build('docs', 'v1', credentials=credentials)
            self.drive_service = build('drive', 'v3', credentials=credentials)
            return True
        except Exception as e:
            logger.error(f"Google authentication error: {str(e)}")
            return False
    
    def create_document(self, title):
        """
        Create a new Google Doc.
        
        Args:
            title (str): Document title
            
        Returns:
            str: Document ID if successful, None otherwise
        """
        if not self.docs_service:
            logger.error("Google Docs service not initialized. Call authenticate() first.")
            return None
            
        try:
            document = self.docs_service.documents().create(
                body={'title': title}
            ).execute()
            return document.get('documentId')
        except HttpError as error:
            logger.error(f"Error creating Google Doc: {error}")
            return None
    
    def apply_mla_format(self, doc_id, metadata):
        """
        Apply MLA formatting to a Google Doc.
        
        Args:
            doc_id (str): Google Doc ID
            metadata (dict): Document metadata
            
        Returns:
            bool: True if successful, False otherwise
        """
        try:
            # Prepare batch update requests
            requests = []
            
            # Set 1-inch margins (72 points)
            requests.append({
                'updateDocumentStyle': {
                    'documentStyle': {
                        'marginTop': {'magnitude': 72, 'unit': 'PT'},
                        'marginBottom': {'magnitude': 72, 'unit': 'PT'},
                        'marginLeft': {'magnitude': 72, 'unit': 'PT'},
                        'marginRight': {'magnitude': 72, 'unit': 'PT'}
                    },
                    'fields': 'marginTop,marginBottom,marginLeft,marginRight'
                }
            })
            
            # Set double spacing for entire document
            requests.append({
                'updateParagraphStyle': {
                    'range': {'startIndex': 1},
                    'paragraphStyle': {
                        'lineSpacing': 2.0
                    },
                    'fields': 'lineSpacing'
                }
            })
            
            # Add student info header
            end_index = 1  # Start at 1 (0 is document beginning)
            
            # Author
            if 'author' in metadata:
                requests.append({
                    'insertText': {
                        'location': {'index': end_index},
                        'text': metadata['author'] + '\n'
                    }
                })
                end_index += len(metadata['author']) + 1
            
            # Instructor
            if 'instructor' in metadata:
                requests.append({
                    'insertText': {
                        'location': {'index': end_index},
                        'text': metadata['instructor'] + '\n'
                    }
                })
                end_index += len(metadata['instructor']) + 1
            
            # Course
            if 'course' in metadata:
                requests.append({
                    'insertText': {
                        'location': {'index': end_index},
                        'text': metadata['course'] + '\n'
                    }
                })
                end_index += len(metadata['course']) + 1
            
            # Date
            if 'date' in metadata:
                requests.append({
                    'insertText': {
                        'location': {'index': end_index},
                        'text': metadata['date'] + '\n\n'
                    }
                })
                end_index += len(metadata['date']) + 2
            
            # Insert title (centered)
            if 'title' in metadata:
                title = metadata['title']
                requests.append({
                    'insertText': {
                        'location': {'index': end_index},
                        'text': title + '\n\n'
                    }
                })
                
                title_start = end_index
                title_end = title_start + len(title)
                
                # Center the title
                requests.append({
                    'updateParagraphStyle': {
                        'range': {
                            'startIndex': title_start,
                            'endIndex': title_end
                        },
                        'paragraphStyle': {
                            'alignment': 'CENTER'
                        },
                        'fields': 'alignment'
                    }
                })
                
                end_index = title_end + 2
            
            # Create a header with page numbers
            requests.append({
                'createHeader': {
                    'type': 'DEFAULT'
                }
            })
            
            # Execute batch updates
            self.docs_service.documents().batchUpdate(
                documentId=doc_id,
                body={'requests': requests}
            ).execute()
            
            # Now get the header sections to add page numbers
            doc = self.docs_service.documents().get(documentId=doc_id).execute()
            
            # Find header section ID
            header_id = None
            for section in doc.get('headers', {}).values():
                header_id = section.get('headerId')
                if header_id:
                    break
            
            if header_id:
                # Add author name and page number to header
                header_requests = [
                    {
                        'insertText': {
                            'location': {'segmentId': header_id, 'index': 0},
                            'text': metadata.get('author', '') + ' '
                        }
                    },
                    {
                        'insertPageNumber': {
                            'location': {'segmentId': header_id, 'index': len(metadata.get('author', '')) + 1}
                        }
                    },
                    {
                        'updateParagraphStyle': {
                            'range': {'segmentId': header_id},
                            'paragraphStyle': {
                                'alignment': 'END'
                            },
                            'fields': 'alignment'
                        }
                    }
                ]
                
                self.docs_service.documents().batchUpdate(
                    documentId=doc_id,
                    body={'requests': header_requests}
                ).execute()
            
            return True
            
        except HttpError as error:
            logger.error(f"Error applying MLA format: {error}")
            return False
    
    def apply_apa_format(self, doc_id, metadata):
        """
        Apply APA formatting to a Google Doc.
        
        Args:
            doc_id (str): Google Doc ID
            metadata (dict): Document metadata
            
        Returns:
            bool: True if successful, False otherwise
        """
        try:
            # Implementation similar to MLA but with APA-specific rules
            # Includes title page, running head, etc.
            
            # Placeholder for full implementation
            requests = []
            
            # Set 1-inch margins
            requests.append({
                'updateDocumentStyle': {
                    'documentStyle': {
                        'marginTop': {'magnitude': 72, 'unit': 'PT'},
                        'marginBottom': {'magnitude': 72, 'unit': 'PT'},
                        'marginLeft': {'magnitude': 72, 'unit': 'PT'},
                        'marginRight': {'magnitude': 72, 'unit': 'PT'}
                    },
                    'fields': 'marginTop,marginBottom,marginLeft,marginRight'
                }
            })
            
            # Execute batch updates
            self.docs_service.documents().batchUpdate(
                documentId=doc_id,
                body={'requests': requests}
            ).execute()
            
            return True
        
        except HttpError as error:
            logger.error(f"Error applying APA format: {error}")
            return False
    
    def apply_chicago_format(self, doc_id, metadata):
        """
        Apply Chicago formatting to a Google Doc.
        
        Args:
            doc_id (str): Google Doc ID
            metadata (dict): Document metadata
            
        Returns:
            bool: True if successful, False otherwise
        """
        try:
            # Implementation similar to MLA but with Chicago-specific rules
            # Includes title page, footnotes, etc.
            
            # Placeholder for full implementation
            requests = []
            
            # Set 1-inch margins
            requests.append({
                'updateDocumentStyle': {
                    'documentStyle': {
                        'marginTop': {'magnitude': 72, 'unit': 'PT'},
                        'marginBottom': {'magnitude': 72, 'unit': 'PT'},
                        'marginLeft': {'magnitude': 72, 'unit': 'PT'},
                        'marginRight': {'magnitude': 72, 'unit': 'PT'}
                    },
                    'fields': 'marginTop,marginBottom,marginLeft,marginRight'
                }
            })
            
            # Execute batch updates
            self.docs_service.documents().batchUpdate(
                documentId=doc_id,
                body={'requests': requests}
            ).execute()
            
            return True
        
        except HttpError as error:
            logger.error(f"Error applying Chicago format: {error}")
            return False
    
    def export_to_docx(self, doc_id, output_path):
        """
        Export Google Doc to DOCX format.
        
        Args:
            doc_id (str): Google Doc ID
            output_path (str): Path to save DOCX file
            
        Returns:
            bool: True if successful, False otherwise
        """
        try:
            if not self.drive_service:
                logger.error("Google Drive service not initialized. Call authenticate() first.")
                return False
                
            # Export the document as DOCX
            response = self.drive_service.files().export(
                fileId=doc_id,
                mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            ).execute()
            
            # Save the file
            with open(output_path, 'wb') as f:
                f.write(response)
                
            return True
            
        except HttpError as error:
            logger.error(f"Error exporting to DOCX: {error}")
            return False
